import streamlit as st
from docx import Document
import random
import re

def extract_mcqs_from_docx(file_path):
    """Extracts MCQs and explanations from a .docx file."""
    document = Document(file_path)
    questions = []
    current_question = {}
    pattern_question = re.compile(r"QUESTION \d+", re.IGNORECASE)
    pattern_answer = re.compile(r"ANSWER: (\w)", re.IGNORECASE)

    for i, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if pattern_question.match(text):
            if current_question:
                questions.append(current_question)
            current_question = {"question": "", "options": [], "answer": "", "explanation": "", "reference": ""}
        elif text.startswith("ANSWER:"):
            match = pattern_answer.match(text)
            if match:
                current_question["answer"] = match.group(1)
            # The explanation is assumed to be in the next paragraph
            if i + 1 < len(document.paragraphs):
                next_paragraph = document.paragraphs[i + 1].text.strip()
                current_question["explanation"] = next_paragraph
        elif text.startswith("Reference:"):
            current_question["reference"] = text.replace("Reference:", "").strip()
        elif current_question and not current_question["answer"]:
            if not current_question["question"]:
                current_question["question"] = text
            else:
                current_question["options"].append(text)
    
    if current_question:
        questions.append(current_question)
    
    return questions

def quiz_app(questions):
    """Runs the quiz app."""
    st.title("MCQ Quiz Game")
    st.markdown("Test your knowledge with the MCQs!")

    # Initialize session state
    if "current_question_index" not in st.session_state:
        st.session_state.current_question_index = 0
        st.session_state.score = 0
        st.session_state.questions_order = random.sample(range(len(questions)), len(questions))  # Random order
        st.session_state.completed = False

    if st.session_state.current_question_index < len(questions):
        question_index = st.session_state.questions_order[st.session_state.current_question_index]
        question = questions[question_index]

        # Display score
        st.sidebar.subheader(f"Current Score: {st.session_state.score}/{st.session_state.current_question_index}")
        
        st.subheader(f"Question {st.session_state.current_question_index + 1}")
        st.write(question["question"])
        
        selected_option = st.radio("Select your answer:", question["options"], key=st.session_state.current_question_index)
        
        if st.button("Submit Answer"):
            correct_option = question["options"][ord(question["answer"]) - ord("A")]
            if selected_option == correct_option:
                st.success("Correct!")
                st.session_state.score += 1
            else:
                st.error(f"Incorrect. Correct answer: {correct_option}")
            
            # Show the explanation
            st.markdown(f"**Explanation:** {question.get('explanation', 'No explanation provided.')}")
            st.markdown(f"**Reference:** {question.get('reference', 'No reference provided.')}")
            
            # Enable "Next Question" button after submitting
            st.session_state.allow_next = True
        
        if st.session_state.get("allow_next", False):
            if st.button("Next Question"):
                st.session_state.current_question_index += 1
                st.session_state.allow_next = False

    if st.session_state.current_question_index == len(questions):
        st.session_state.completed = True

    if st.session_state.completed:
        st.balloons()
        st.subheader("Quiz Completed!")
        st.write(f"Your final score: {st.session_state.score} / {len(questions)}")
        st.button("Restart Quiz", on_click=reset_quiz)

def reset_quiz():
    """Resets the quiz."""
    st.session_state.current_question_index = 0
    st.session_state.score = 0
    st.session_state.questions_order = random.sample(range(len(st.session_state.questions)), len(st.session_state.questions))
    st.session_state.completed = False
    st.session_state.allow_next = False

def main():
    st.title("MCQ Quiz Generator")
    st.markdown("Answer the questions to test your knowledge.")

    # Load the .docx file from the code repo
    file_path = "mcqs.docx"  # Replace with the actual file path in your repo
    questions = extract_mcqs_from_docx(file_path)

    # Store questions in session state
    if "questions" not in st.session_state:
        st.session_state.questions = questions

    if questions:
        quiz_app(questions)
    else:
        st.error("No valid questions found in the .docx file.")

if __name__ == "__main__":
    main()